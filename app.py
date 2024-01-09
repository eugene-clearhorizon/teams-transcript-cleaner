from flask import *
from werkzeug.utils import secure_filename
from werkzeug.datastructures import  FileStorage
import os
from docx import Document
import unicodedata
import pandas as pd
import time
from datetime import datetime, timedelta


app = Flask(__name__)
# Configure file uploads
# Update the allowed extensions to include '.docx'
ALLOWED_EXTENSIONS = {'docx'}
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16 MB limit
app.config['CLEANED_UPLOADS_FOLDER'] = 'cleaned_uploads'

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def cleanup_old_files(folder, days=7):
    now = datetime.now()
    for filename in os.listdir(folder):
        file_path = os.path.join(folder, filename)
        if os.path.isfile(file_path):
            # Get the last modification time of the file
            modified_time = datetime.fromtimestamp(os.path.getmtime(file_path))
            print(modified_time)
            # Calculate the age of the file
            age = now - modified_time
            print(age)
            # If the file is older than the specified days, delete it
            if age > timedelta(days=days):
                os.remove(file_path)

def clean_transcripts(file_path):
    start_time = time.time()             
    
    cleaned_uploads_dir = 'cleaned_uploads'
    if not os.path.exists(cleaned_uploads_dir):
        os.makedirs(cleaned_uploads_dir)

    ### imports a word document and returns it as a single string
    # Create a Document object from the file
    doc = Document(file_path)
    # Initialize an empty list to store the normalized paragraphs
    text = []    
    for p in doc.paragraphs:
        # Normalize the text of the paragraph to remove unwanted unicode characters
        result = unicodedata.normalize('NFKD', p.text)
        text.append(f"{result}\n")
    # Join all the items in the list into a single string
    str_text = ''.join(text)
    # Split the string into a list of lines and remove the line break characters
    final_string = str_text.split('\n')

    # a set that will remove lines that only contain these unnecessary 'filler' words
    unwanted_words = {'','Uh-huh.', 'Yeah. And so.' , 'OK. Yeah, yeah. ' , 'Umm.', 'Yeah.', 'Yeah, yeah.', 'Awesome.', 
                        'OK. Yep.', 'OK.', 'Right.', 'Right?', 'OK. Yeah.','So.','Uh.',
                        'Hmm.', 'Hmm yeah.', 'Yeah, cool.', 'Ohh.', 'Um', 'Um?', 'Umm?', 'Cool.', 'Mm-hmm.', 'Mm hmm.', 'Huh.',
                        'Ohh uh-huh.', 'Yeah. Wow.', 'Ohh wow wow.'}

    # Remove all lines containing '-->', these are the timestamps
    list_one = [i for i in final_string if '-->' not in i]   

    # Remove all lines consisting solely of the unwanted words
    list_two = [i for i in list_one if i not in unwanted_words]          

    # Get the Names of the participants by finding most common values in the list
    n = 2
    df = pd.DataFrame(list_two)    
    names = set(df[0].value_counts()[:n].index.tolist())
    print(names)

    # Remove lines where a participants name appears but what they said has been deleted. After doing this we will
    # only see a pattern of name, response, name, response etc
    list_three = []
    for n, i in enumerate(list_two):                   
        if i not in names:
            list_three.append(list_two[n-1])
            list_three.append(i)

    # when a partcipants speaks multiple times in a row, remove subsequent lines with their name so the name
    # only appears once each time someone speaks
    list_four = []
    for n, i in enumerate(list_three):        
        if n < 2:
            list_four.append(i)
        elif i != list_three[n-2]:
            list_four.append(i)

    ### Compiles and saves the clean transcript into a new word document
    # create a new document
    new_doc = Document()
    for i in list_four:
        # add the cleaned transcript into a new paragraph
        if i in names:
            paragraph = new_doc.add_paragraph()
            paragraph.add_run(f"{i}: ").bold = True
        else:
            paragraph.add_run(f"{i} ")
    # save the document with a given file name
    cleaned_filename = f"cleaned_{os.path.basename(file_path)}"
    new_doc.save(os.path.join('cleaned_uploads', cleaned_filename))

    os.remove(file_path)        
    
    print(f"--- {(time.time() - start_time)} seconds ---")

    return names, cleaned_filename



@app.route('/')
def index():
    # Call the cleanup function before rendering the template    
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_files():
    try:
        # Call the cleanup function before processing the uploaded files
        cleanup_old_files(app.config['CLEANED_UPLOADS_FOLDER'])
        # Check if the post request has the file part
        if 'file' not in request.files:
            return 'No file part'

        files = request.files.getlist('file')
        num_files_uploaded = 0
        total_time = 0
        file_info_list = []

        # Iterate through each uploaded file
        for file in files:
            if file.filename == '':
                continue

            if file and allowed_file(file.filename):
                # Ensure that the 'uploads' directory exists
                uploads_dir = 'uploads'
                if not os.path.exists(uploads_dir):
                    os.makedirs(uploads_dir)

                # Save the uploaded file to the 'uploads' directory
                filename = secure_filename(file.filename)
                file_path = os.path.join(uploads_dir, filename)
                file.save(file_path)

                # Measure the time taken to clean transcripts for each file                
                start_time = time.time()
                names, cleaned_filename = clean_transcripts(file_path)
                processing_time = time.time() - start_time

                # Move the cleaned file to the 'cleaned_uploads' directory
                cleaned_path = os.path.join(app.config['CLEANED_UPLOADS_FOLDER'], cleaned_filename)
                
                # Collect information about each file
                file_info = {
                    'filename': cleaned_filename,
                    'names': names,
                    'processing_time': processing_time
                }
                file_info_list.append(file_info)

                total_time += processing_time
                num_files_uploaded += 1
                print(os.path.join(current_app.root_path, app.config['CLEANED_UPLOADS_FOLDER']))

        return render_template('upload_summary.html', num_files=num_files_uploaded, total_time=total_time, file_info_list=file_info_list)
    except Exception as e:
        return f'An error occurred: {str(e)}'

@app.route('/cleaned_uploads/<path:filename>', methods=['GET', 'POST'])
def download_cleaned(filename):
    cleaned_uploads = os.path.join(current_app.root_path, app.config['CLEANED_UPLOADS_FOLDER'])
    print(os.path.join(current_app.root_path, app.config['CLEANED_UPLOADS_FOLDER']))
    return send_from_directory(cleaned_uploads, filename)

if __name__ == '__main__':
    app.run(debug=True)